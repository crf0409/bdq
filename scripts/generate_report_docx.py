# -*- coding: utf-8 -*-
"""
生成实验报告图表汇总docx文档, 并与所有图片一起打包为zip
"""

import os
import sys
import zipfile
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_ROOT = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..')
os.chdir(PROJECT_ROOT)

FIGURES_DIR = Path('results/report_figures')
OUTPUT_DIR = Path('results')

# 图表清单: (文件名, 图号标题)
FIGURES = [
    ("fig3_5_training_curve", "图3-5 / 图5-7 SAC训练奖励曲线"),
    ("fig5_1_pv_60s", "图5-1 光伏出力快速变化(60s)"),
    ("fig5_2_voltage_60s", "图5-2 60s内系统节点电压分布"),
    ("fig5_3_droop_deadband", "图5-3 带死区的传统Q-V下垂特性"),
    ("fig5_3_pv_6node", "图5-3补充 6节点光伏出力曲线"),
    ("fig5_4_droop_improved", "图5-4 改进下垂特性(可调电压截距)"),
    ("fig5_4_dynamic_load", "图5-4补充 动态负荷曲线"),
    ("fig5_5_voltage_prob", "图5-5 电压概率分布直方图(无电压控制 vs 所提方法)"),
    ("fig5_6_loss_bar", "图5-6 各时段网络损耗分布柱状图"),
    ("fig5_8_no_ctrl_v12", "图5-8 无电压控制 t=12:00节点电压分布"),
    ("fig5_9_deadband_v12", "图5-9 传统死区下垂 t=12:00节点电压分布"),
    ("fig5_10_improved_v12", "图5-10 改进下垂控制 t=12:00节点电压分布"),
    ("fig5_11_sac_v12", "图5-11 SAC改进下垂 t=12:00节点电压分布"),
    ("fig5_12_four_strategy_v12", "图5-12 t=12:00四种策略节点电压分布对比"),
    ("fig5_13_node17_voltage", "图5-13 节点17全天电压变化对比"),
    ("fig5_14_max_voltage", "图5-14 全网最高电压全天变化对比"),
    ("fig5_15_min_voltage", "图5-15 全网最低电压全天变化对比"),
    ("fig5_16_loss_no_ctrl", "图5-16 无电压控制 24小时网损曲线"),
    ("fig5_17_loss_deadband", "图5-17 传统死区下垂 24小时网损曲线"),
    ("fig5_18_loss_improved", "图5-18 改进下垂控制 24小时网损曲线"),
    ("fig5_19_loss_sac", "图5-19 SAC改进下垂 24小时网损曲线"),
    ("fig5_20_loss_period_bar", "图5-20 各时段网络损耗对比"),
    ("fig5_21_loss_combined", "图5-21 四种策略24小时网损曲线对比"),
    ("fig5_22_loss_total_bar", "图5-22 四种策略24小时总网损对比"),
    ("fig5_23_soc", "图5-23 储能SOC全天变化曲线"),
    ("fig5_24_q_output", "图5-24 t=12:00各PV逆变器无功输出对比"),
    ("fig5_25_voltage_prob_4strategy", "图5-25 四种策略电压概率分布直方图"),
]


def generate_docx():
    """生成包含所有图表的docx文档"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # 标题
    title = doc.add_heading('实验报告图表汇总', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')  # 空行

    for fname, caption in FIGURES:
        img_path = FIGURES_DIR / f'{fname}.png'
        if not img_path.exists():
            print(f"  警告: 未找到 {img_path}, 跳过")
            continue

        # 插入图片, 宽度适应A4页面
        doc.add_picture(str(img_path), width=Cm(15))

        # 图片居中
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加图题
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_para.add_run(caption)
        run.font.size = Pt(10)
        run.font.name = '宋体'

        # 图间空行
        doc.add_paragraph('')

    docx_path = OUTPUT_DIR / '实验报告图表汇总.docx'
    doc.save(str(docx_path))
    print(f"已生成docx: {docx_path}")
    return docx_path


def create_zip(docx_path):
    """将docx和所有png图片打包为zip"""
    zip_path = OUTPUT_DIR / '实验报告图表.zip'

    with zipfile.ZipFile(str(zip_path), 'w', zipfile.ZIP_DEFLATED) as zf:
        # 添加docx
        zf.write(str(docx_path), docx_path.name)
        print(f"  添加: {docx_path.name}")

        # 添加所有png图片
        for fname, _ in FIGURES:
            img_path = FIGURES_DIR / f'{fname}.png'
            if img_path.exists():
                arcname = f'report_figures/{fname}.png'
                zf.write(str(img_path), arcname)

        png_count = sum(1 for f, _ in FIGURES if (FIGURES_DIR / f'{f}.png').exists())
        print(f"  添加: {png_count} 张PNG图片")

    size_mb = zip_path.stat().st_size / 1024 / 1024
    print(f"已生成zip: {zip_path} ({size_mb:.1f} MB)")
    return zip_path


def main():
    print("=" * 50)
    print("生成实验报告docx + zip打包")
    print("=" * 50)

    # 检查图片是否存在
    missing = [f for f, _ in FIGURES if not (FIGURES_DIR / f'{f}.png').exists()]
    if missing:
        print(f"警告: 缺少 {len(missing)} 张图片: {missing}")
        print("请先运行: python scripts/generate_report_figures.py")

    # 生成docx
    print("\n生成docx文档...")
    docx_path = generate_docx()

    # 打包zip
    print("\n打包zip...")
    zip_path = create_zip(docx_path)

    print("\n完成!")
    print(f"  docx: {docx_path.absolute()}")
    print(f"  zip:  {zip_path.absolute()}")


if __name__ == "__main__":
    main()
